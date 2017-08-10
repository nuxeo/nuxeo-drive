# coding: utf-8
import os
import random
import string

from PIL import Image
from docx import Document

words = dict()


def load_dictionary(filename):
    with open(filename, "r") as f:
        for word in f.readlines():
            word = word.strip()
            l = len(word)
            if not l in words:
                words[l] = list()
            words[l].append(word)
    for key in words.keys():
        print "%d words of length %d" % (len(words[key]), key)


def id_generator(size=6, chars=string.ascii_letters + string.digits):
    if (not size in words):
        # Generate random as no words exists
        return ''.join(random.choice(chars) for _ in range(size))
    else:
        return random.choice(words[size])


def generate_random_string(size=None):
    if size is None:
        size = random.randint(5, 30)
    res = ""
    while (size > 0):
        if size > 20:
            word_size = random.randint(4, 20)
            res = res + " " + id_generator(word_size)
            word_size = word_size + 1
        elif size > 2:
            word_size = random.randint(1, size - 1)
            res = res + " " + id_generator(word_size)
            word_size = word_size + 1
        else:
            word_size = size
            res = res + id_generator(size)
        size = size - word_size
    return res


def generate_random_docx(filename, paraph_size=30):
    document = Document()

    size = random.randint(5, paraph_size)
    document.add_heading(generate_random_string(size), 0)

    size = random.randint(5, paraph_size)
    p = document.add_paragraph(
                               generate_random_string(size))
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    size = random.randint(5, paraph_size)
    document.add_heading(generate_random_string(size), level=1)
    size = random.randint(5, paraph_size)
    document.add_paragraph(generate_random_string(size), style='IntenseQuote')

    for _ in range(random.randint(1, 10)):
        size = random.randint(5, paraph_size)
        document.add_paragraph(
                        generate_random_string(size), style='ListBullet')

    for _ in range(random.randint(1, 10)):
        size = random.randint(5, paraph_size)
        document.add_paragraph(
                        generate_random_string(size), style='ListNumber')

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for i in range(random.randint(1, 10)):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = generate_random_string(8)
        row_cells[2].text = generate_random_string(30)

    document.add_page_break()

    document.save(filename)


def generate_random_jpg(filename, size):
    import numpy
    a = numpy.random.rand(size, size, 3) * 255
    im_out = Image.fromarray(a.astype('uint8')).convert('RGBA')
    im_out.save(filename)
    pass


def generate_random_txt(filename, size):
    with open(filename, 'w') as f:
        for _ in range(1024):
            f.write(generate_random_string(size))


def generate_filename(path, extension=""):
    filename = os.path.join(path,
                        id_generator(random.randint(4, 26)) + extension)
    while os.path.exists(filename):
        filename = os.path.join(path,
                        id_generator(random.randint(4, 26)) + extension)
    return filename


def generate_folder_files(path, files_number):
    if files_number == 0:
        return
    if files_number < 3:
        for _ in range(files_number):
            generate_random_jpg(generate_filename(path, ".jpg"), random.randint(1000,3000))
        return
    # Calculate images, docx, jpg
    jpeg = random.randint(1, files_number - 2)
    files_number = files_number - jpeg
    docx = random.randint(1, files_number - 1)
    txt = files_number - docx
    for _ in range(jpeg):
        generate_random_jpg(generate_filename(path, ".jpg"), random.randint(1000,3000))
    for _ in range(docx):
        generate_random_docx(generate_filename(path, ".docx"),
                                random.randint(30, 600))
    for _ in range(txt):
        generate_random_txt(generate_filename(path, ".txt"),
                            random.randint(10, 4096))


def generate_files(path, folder_number, files_number, size_min=10, size_max=3000,
                    depth=5, additional_files=None):
    import math
    if additional_files is None:
        additional_files = []
    print "Files number: %d\tFolder number: %d\tDepth: %d" % (files_number, folder_number, depth)
    if depth > 0:
        # Get the max number of folders
        if folder_number != 0:
            local_files = random.randint(1, math.ceil(float(files_number) / float(folder_number)) * 2)
        else:
            local_files = files_number
        if local_files > files_number:
            local_files = files_number
        if depth > 1:
            folder_repart = math.ceil(math.pow(folder_number, 1.0 / depth))
            local_folders = random.randint(1, folder_repart)
        else:
            local_folders = folder_number
        print "Create %d folders and %d files in %s" % (local_folders, local_files, path)
        folder_number = folder_number - local_folders
        files_number = files_number - local_files
        for i in range(local_folders):
            name = generate_filename(path)
            os.mkdir(name)
            if i == local_folders - 1:
                print "Take the rest"
                nxt_folder = folder_number
                nxt_files = files_number
            else:
                if folder_number > 0:
                    nxt_folder = random.randint(1, math.ceil(float(folder_number) / float(local_folders-i)))
                else:
                    nxt_folder = 0
                if files_number > 0:
                    nxt_files = random.randint(1, math.ceil(float(files_number) / float(local_folders-i)))
                else:
                    nxt_files = 0
            folder_number = folder_number - nxt_folder
            files_number = files_number - nxt_files
            generate_files(name, nxt_folder, nxt_files, size_min, size_max, depth-1, additional_files)
        generate_folder_files(path, local_files)
    else:
        print "Create %d files in %s" % (files_number, path)
        generate_folder_files(path, files_number)


if __name__ == '__main__':
    load_dictionary("/usr/share/dict/web2a")
    path = 'benchmark_files'
    if not os.path.exists(path):
        os.mkdir(path)
    generate_files(path, 100, 800, depth=2)
